# -*- coding: utf-8 -*-
"""Сборка одного Word-файла из всех страниц книги (pages.json)."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Дискурс.docx"
TOKEN = re.compile(
    r"!\[\[([^\]|]+)(?:\|([^\]]*))?\]\]"  # Obsidian image
    r"|!\[([^\]]*)\]\(([^)]+)\)"  # markdown image
    r"|##\s+(.+)"  # heading (whole-line, checked separately)
)
WIKI_OR_MD = re.compile(
    r"!\[\[([^\]|]+)(?:\|([^\]]*))?\]\]|!\[([^\]]*)\]\(([^)]+)\)"
)


INCLUDE_RE = re.compile(r"^(\d+(?:\.\d+)*)-(\d+(?:\.\d+)*)")


def parse_include(path: str, explicit: str | None = None) -> str | None:
    if explicit:
        return explicit
    name = Path(path).name.replace(".md", "")
    match = INCLUDE_RE.match(name)
    return match.group(2) if match else None


def title_from_md(path: str) -> str:
    name = Path(path).name.replace(".md", "")
    title = re.sub(r"^\d+(?:\.\d+)*(?:-\d+(?:\.\d+)*)?\.?\s*", "", name).strip()
    return title or name.rstrip(".")


def find_image(name: str) -> Path | None:
    name = name.strip()
    direct = ROOT / name
    if direct.is_file():
        return direct
    for path in ROOT.rglob(name):
        if path.is_file():
            return path
    return None


def set_run_font(run, size_pt: float = 12, bold: bool = False, italic: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, 12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_image(doc: Document, path: Path, width_hint: str | None = None):
    width = Cm(14)
    if width_hint and str(width_hint).isdigit():
        width = Cm(min(16, max(4, int(width_hint) / 40)))
    try:
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as exc:
        add_text_paragraph(doc, f"[Не удалось вставить изображение: {path.name} — {exc}]")


def insert_image_ref(doc: Document, filename: str, width_hint: str | None = None):
    img = find_image(filename)
    if img:
        add_image(doc, img, width_hint)
    else:
        add_text_paragraph(doc, f"[Нет файла изображения: {filename}]")


def flush_paragraph(doc: Document, lines: list[str]):
    text = "\n".join(lines).strip()
    if text:
        add_text_paragraph(doc, text)
    lines.clear()


def render_mixed_line(doc: Document, line: str):
    """Line may mix text and images."""
    pos = 0
    text_bits: list[str] = []

    def flush_text():
        joined = " ".join(t for t in text_bits if t).strip()
        text_bits.clear()
        if joined:
            add_text_paragraph(doc, joined)

    for match in WIKI_OR_MD.finditer(line):
        before = line[pos : match.start()].strip()
        if before:
            text_bits.append(before)
        flush_text()
        if match.group(1) is not None:
            insert_image_ref(doc, match.group(1), match.group(2))
        else:
            insert_image_ref(doc, Path(match.group(4)).name)
        pos = match.end()

    after = line[pos:].strip()
    if after:
        text_bits.append(after)
    flush_text()


def render_page_body(doc: Document, text: str):
    lines_buf: list[str] = []
    for raw in text.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()

        if not line.strip():
            flush_paragraph(doc, lines_buf)
            continue

        if line.strip().startswith("## "):
            flush_paragraph(doc, lines_buf)
            p = doc.add_heading(line.strip()[3:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, 14, bold=True)
            continue

        if WIKI_OR_MD.search(line):
            flush_paragraph(doc, lines_buf)
            render_mixed_line(doc, line)
            continue

        lines_buf.append(line.strip())

    flush_paragraph(doc, lines_buf)


def build():
    pages = json.loads((ROOT / "pages.json").read_text(encoding="utf-8"))["pages"]
    by_id = {p["id"]: p for p in pages}
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_heading("Дискурс", level=0)
    for run in title.runs:
        set_run_font(run, 28, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Читать и слушать")
    set_run_font(run, 14, italic=True)
    run.font.color.rgb = RGBColor(0x6B, 0x5F, 0x4D)

    pages = [
        p
        for p in pages
        if not str(p.get("section", "")).rstrip().endswith("@")
        and not Path(p.get("md", "")).stem.rstrip().endswith("@")
    ]

    current_section = None
    for i, page in enumerate(pages):
        if page["section"] != current_section:
            current_section = page["section"]
            h = doc.add_heading(current_section, level=1)
            for run in h.runs:
                set_run_font(run, 18, bold=True)
                run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)

        include_id = parse_include(page["md"], page.get("include"))
        source = by_id.get(include_id) if include_id else None
        display_title = (
            title_from_md(source["md"]) if source else title_from_md(page["md"])
        )
        page_title = f"{page['id']}. {display_title}".strip()
        h2 = doc.add_heading(page_title, level=2)
        for run in h2.runs:
            set_run_font(run, 14, bold=True)

        # Include pages: never use own .md body — only the source page
        if include_id:
            if not source:
                add_text_paragraph(
                    doc,
                    f"[Вставка {page['id']}←{include_id}: источник не найден; "
                    f"собственный текст игнорируется]",
                )
            else:
                md_path = ROOT / source["md"]
                if not md_path.exists():
                    add_text_paragraph(doc, f"[Файл не найден: {source['md']}]")
                else:
                    render_page_body(doc, md_path.read_text(encoding="utf-8"))
        else:
            md_path = ROOT / page["md"]
            if not md_path.exists():
                add_text_paragraph(doc, f"[Файл не найден: {page['md']}]")
            else:
                render_page_body(doc, md_path.read_text(encoding="utf-8"))

        if i < len(pages) - 1:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

    doc.save(OUTPUT)
    print(f"OK: {OUTPUT}")
    print(f"Страниц: {len(pages)}")


if __name__ == "__main__":
    build()
